import streamlit as st
import pandas as pd
import pytesseract
from pdf2image import convert_from_bytes
import io
import numpy as np
import re
from decimal import Decimal, ROUND_HALF_UP
import google.generativeai as genai
from docx import Document

# ------------------------------------------------------------------
# CONFIGURATION & SECURITY
# ------------------------------------------------------------------
st.set_page_config(page_title="Secure AI Data Tool", layout="wide")

# --- SECURITY: PASSWORD PROTECTION ---
def check_password():
    """Returns `True` if the user had the correct password."""
    if "app_password" not in st.secrets:
        st.error("❌ No password set in secrets.toml! Please configure it.")
        return False

    def password_entered():
        if st.session_state["password"] == st.secrets["app_password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("🔒 App Password:", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("🔒 App Password:", type="password", on_change=password_entered, key="password")
        st.error("😕 Password incorrect")
        return False
    else:
        return True

if not check_password():
    st.stop()

# ------------------------------------------------------------------
# MAIN APP
# ------------------------------------------------------------------
st.title("🧠 Secure Smart Scan & Analyze Tool")
tab1, tab2, tab3 = st.tabs(["📄 Scan PDF to Excel", "🤖 Analyze Excel", "📝 Generate Reports (Docs)"])

# ------------------------------------------------------------------
# HELPER FUNCTIONS
# ------------------------------------------------------------------
def strict_invoice_round(x):
    """Rounds to exactly 2 decimal places (Standard Rounding)"""
    try:
        if pd.isna(x) or str(x).strip() == "": return 0.0
        d = Decimal(str(x))
        return float(d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))
    except: return 0.0

def process_layout_preserving(image, clustering_sensitivity=15):
    """Core OCR Logic for Tables"""
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    df = pd.DataFrame(data)
    df = df[df['text'].str.strip() != '']
    if df.empty: return pd.DataFrame()
    
    df['row_id'] = (df['top'] / 15).round().astype(int) 
    all_lefts = df['left'].sort_values().unique()
    col_definitions = [] 
    for x in all_lefts:
        found = False
        for i, c in enumerate(col_definitions):
            if abs(x - c) < clustering_sensitivity:
                col_definitions[i] = (c + x) / 2
                found = True; break
        if not found: col_definitions.append(x)
    col_definitions.sort()
    
    if not col_definitions: return pd.DataFrame()

    df['col_idx'] = df['left'].apply(lambda x: np.argmin([abs(x - c) for c in col_definitions]))
    
    grid = [['' for _ in range(len(col_definitions))] for _ in range(len(df['row_id'].unique()))]
    row_map = {rid: i for i, rid in enumerate(sorted(df['row_id'].unique()))}
    for _, row in df.iterrows():
        grid[row_map[row['row_id']]][row['col_idx']] = (grid[row_map[row['row_id']]][row['col_idx']] + " " + row['text']).strip()
        
    final = pd.DataFrame(grid)
    final = final.loc[:, (final != '').any(axis=0)]
    final.columns = [f"Col_{i+1}" for i in range(final.shape[1])]
    return final

# ------------------------------------------------------------------
# TAB 1: SCAN PDF TO EXCEL
# ------------------------------------------------------------------
with tab1:
    st.header("1. Scan PDF to Excel")
    if 'scan_df' not in st.session_state: st.session_state.scan_df = None
    uploaded_pdf = st.file_uploader("Upload Scanned PDF", type=["pdf"], key="pdf1")
    
    if uploaded_pdf:
        with st.expander("Alignment Settings"):
            sens = st.slider("Sensitivity", 5, 100, 25)
        try:
            images = convert_from_bytes(uploaded_pdf.read())
            st.session_state.scan_df = process_layout_preserving(images[0], sens)
            st.success("Scan Processed!")
            edited_scan = st.data_editor(st.session_state.scan_df, num_rows="dynamic", use_container_width=True)
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer: edited_scan.to_excel(writer, index=False, header=False)
            st.download_button("📥 Download Excel", output.getvalue(), "scan.xlsx")
        except Exception as e: st.error(f"Error: {e}")

# ------------------------------------------------------------------
# TAB 2: ANALYZE EXCEL
# ------------------------------------------------------------------
with tab2:
    st.header("2. Ask the App to Calculate")
    uploaded_excel = st.file_uploader("Upload Excel File", type=["xlsx", "xls"], key="xls1")
    
    if uploaded_excel:
        st.info("Check preview. Adjust Header Row if needed.")
        header_row_idx = st.number_input("Header Row Number", min_value=0, max_value=20, value=0)
        df_excel = pd.read_excel(uploaded_excel, header=header_row_idx)
        
        with st.expander("👀 View Data Preview", expanded=True):
            st.dataframe(df_excel.head().round(2))
            st.caption(f"Valid Columns: {', '.join(list(df_excel.columns))}")
        
        st.divider()
        user_query = st.text_input("Direction:", placeholder="e.g. Sort by Date. Sum all columns by Class.")
        run_calc = st.button("🚀 Run Analysis", type="primary")

        if run_calc and user_query:
            steps = re.split(r'[.;]| then | after that | and ', user_query, flags=re.IGNORECASE)
            steps = [s.strip() for s in steps if s.strip()]
            current_df = df_excel.copy()
            final_result = None
            log_messages = []
            
            try:
                for step in steps:
                    step_lower = step.lower()
                    col_map = {c.lower(): c for c in current_df.columns}
                    sorted_cols = sorted(col_map.keys(), key=len, reverse=True)
                    found_cols = [col_map[c] for c in sorted_cols if c in step_lower]
                    
                    op = None
                    if 'sort' in step_lower: op = 'sort'
                    elif any(x in step_lower for x in ['sum', 'total', 'add', 'subtotal']): op = 'sum'
                    elif any(x in step_lower for x in ['avg', 'average']): op = 'mean'
                    elif any(x in step_lower for x in ['count']): op = 'count'
                    
                    if op == 'sort' and found_cols:
                        target = found_cols[0]
                        ascending = False if 'desc' in step_lower else True
                        current_df = current_df.sort_values(by=target, ascending=ascending)
                        log_messages.append(f"✅ Sorted by **{target}**")
                        final_result = current_df 
                    elif op in ['sum', 'mean', 'count']:
                        group_col = None
                        if 'by' in step_lower:
                            parts = step_lower.split('by')
                            cands = [c for c in found_cols if c.lower() in parts[1]]
                            if cands: group_col = cands[0]
                        
                        if group_col:
                            val_cols = [c for c in found_cols if c != group_col]
                            if 'all' in step_lower or 'each' in step_lower or not val_cols:
                                target_cols = current_df.select_dtypes(include=[np.number]).columns.tolist()
                                if group_col in target_cols: target_cols.remove(group_col)
                            else: target_cols = val_cols

                            if target_cols:
                                for c in target_cols:
                                     current_df[c] = pd.to_numeric(current_df[c].astype(str).str.replace(r'[^\d\.\-]', '', regex=True), errors='coerce')
                                res = current_df.groupby(group_col)[target_cols].agg(op)
                                res_df = res.reset_index()
                                res_df.columns = [group_col] + [f"{op.title()} of {c}" for c in target_cols]
                                final_result = res_df
                                log_messages.append(f"✅ Calculated **{op}** for **{len(target_cols)} columns** by **{group_col}**")
                        else:
                            if found_cols:
                                target = found_cols[0]
                                current_df[target] = pd.to_numeric(current_df[target].astype(str).str.replace(r'[^\d\.\-]', '', regex=True), errors='coerce')
                                val = current_df[target].agg(op)
                                final_result = pd.DataFrame({f"{op.title()} of {target}": [val]})
                                log_messages.append(f"✅ Calculated total **{op}** of **{target}**")
            except Exception as e: st.error(f"Error: {e}")

            for msg in log_messages: st.write(msg)
            if final_result is not None:
                res_numeric = final_result.select_dtypes(include=['float', 'float64']).columns
                disp = final_result.copy()
                disp[res_numeric] = disp[res_numeric].round(2)
                st.dataframe(disp, use_container_width=True)
                out = io.BytesIO()
                with pd.ExcelWriter(out, engine='openpyxl') as writer: disp.to_excel(writer, index=False)
                st.download_button("📥 Download Result", out.getvalue(), "result.xlsx")

# ------------------------------------------------------------------
# TAB 3: GENERATE REPORTS (AUTO-DETECT MODEL)
# ------------------------------------------------------------------
with tab3:
    st.header("3. Generate Word/Google Docs")
    st.markdown("This uses the API Key stored in your `secrets.toml`.")
    
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
        genai.configure(api_key=api_key)
    else:
        st.error("❌ No API Key found in secrets.toml.")
        st.stop()

    uploaded_doc_pdf = st.file_uploader("Upload PDF for Report", type=["pdf"], key="pdf_docs")
    
    if uploaded_doc_pdf:
        images = convert_from_bytes(uploaded_doc_pdf.read())
        raw_text = ""
        for img in images:
            raw_text += pytesseract.image_to_string(img) + "\n"
        st.success("PDF Loaded!")
        
        # --- MODEL SELECTOR ---
        st.subheader("AI Model Selection")
        
        # 1. Try to fetch available models dynamically from the user's key
        try:
            available_models = [
                m.name for m in genai.list_models() 
                if 'generateContent' in m.supported_generation_methods
            ]
            # Clean up names (remove 'models/')
            clean_models = [m.replace('models/', '') for m in available_models]
            
            # Prioritize Gemini 2.5, then 2.0, then 1.5
            # We sort them to show the newest first
            clean_models.sort(reverse=True)
            
        except Exception as e:
            # Fallback if list_models fails
            clean_models = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash"]
        
        # 2. Let user pick (Default to first in list, which should be newest)
        col_model, col_style = st.columns(2)
        with col_model:
            selected_model = st.selectbox("Select AI Model:", clean_models, index=0)
        
        with col_style:
            template_choice = st.radio("Style:", ["Data Extraction (Table)", "Research Summary (Report)"])
        
        if "Data" in template_choice:
            default_prompt = "Extract data into a Markdown table. No chat text, just the table."
        else:
            default_prompt = "Summarize this into a formal report with headers and bullet points."

        user_prompt = st.text_area("Instructions:", value=default_prompt, height=100)
        
        if st.button("✨ Generate"):
            try:
                with st.spinner(f"Using {selected_model}..."):
                    model = genai.GenerativeModel(selected_model)
                    response = model.generate_content(f"{user_prompt}\n\nDocument:\n{raw_text}")
                    generated_text = response.text
                    
                    st.markdown(generated_text)
                    
                    doc = Document()
                    doc.add_heading('Generated Report', 0)
                    for line in generated_text.split('\n'):
                        clean = line.strip()
                        if clean.startswith('## '): doc.add_heading(clean[3:], level=2)
                        elif clean.startswith('# '): doc.add_heading(clean[2:], level=1)
                        elif clean.startswith('* ') or clean.startswith('- '): doc.add_paragraph(clean[2:], style='List Bullet')
                        else: doc.add_paragraph(clean)
                    
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button("📥 Download .docx", buf, "report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"Error: {e}")
